from io import BytesIO
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document


class RegistrationData(BaseModel):
    beneficiary1Title: str = ""
    beneficiary1LastName: str = ""
    beneficiary1FirstName: str = ""
    beneficiary1BirthDate: str = ""
    beneficiary2Title: str = ""
    beneficiary2LastName: str = ""
    beneficiary2FirstName: str = ""
    beneficiary2BirthDate: str = ""
    socialSecurityNumber: str = ""
    addressLine1: str = ""
    addressLine2: str = ""
    postalCode: str = ""
    floor: str = ""
    mobilePhone: str = ""
    homePhone: str = ""
    email: str = ""
    hasLegalRepresentative: bool = False
    legalRepLastName: str = ""
    legalRepFirstName: str = ""
    legalRepAddressLine1: str = ""
    legalRepAddressLine2: str = ""
    legalRepPhone: str = ""
    legalRepEmail: str = ""
    emergency1LastName: str = ""
    emergency1FirstName: str = ""
    emergency1AddressLine1: str = ""
    emergency1AddressLine2: str = ""
    emergency1Phone: str = ""
    emergency1Email: str = ""
    emergency1Relation: str = ""
    emergency2LastName: str = ""
    emergency2FirstName: str = ""
    emergency2AddressLine1: str = ""
    emergency2AddressLine2: str = ""
    emergency2Phone: str = ""
    emergency2Email: str = ""
    emergency2Relation: str = ""
    mobilityType: str = ""
    aidWalker: bool = False
    aidTransferChair: bool = False
    aidTripodCane: bool = False
    aidQuadripodCane: bool = False
    aidSimpleCane: bool = False
    aidCrutch: bool = False
    docResidenceProof: bool = False
    docIdentityCard: bool = False
    docVitaleCard: bool = False
    docRetirementOrASPA: bool = False
    docAPA: bool = False
    docPCH: bool = False
    docCMIPriorityOrInvalidity: bool = False
    docMedicalCertificate: bool = False
    engagementTransportRules: bool = False
    attestAccuracy: bool = False
    signatureDate: str = ""
    additionalNotes: str = ""


app = FastAPI(title="LaReinette Docx API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "templates" / "Fiche-Inscription-Template.docx"


def safe(value: str) -> str:
    txt = (value or "").strip()
    return txt if txt else "-"


def yes_no(value: bool) -> str:
    return "Oui" if value else "Non"


def add_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> bool:
    text = paragraph.text
    replaced = False
    for key, value in replacements.items():
        if key in text:
            text = text.replace(key, value)
            replaced = True
    if replaced:
        paragraph.text = text
    return replaced


def replace_in_doc(doc: Document, replacements: dict[str, str]) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, replacements):
            count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, replacements):
                        count += 1

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if replace_in_paragraph(paragraph, replacements):
                count += 1
        for paragraph in section.footer.paragraphs:
            if replace_in_paragraph(paragraph, replacements):
                count += 1
    return count


@app.get("/")
def home():
    return {"message": "Bienvenue sur l'API La Reinette", "status": "running"}


@app.get("/health")
def health():
    return {"ok": True}


@app.post("/generate-docx")
def generate_docx(data: RegistrationData):
    try:
        replacements = {
            "{{date_formulaire}}": datetime.now().strftime("%d/%m/%Y"),
            "{{beneficiary1Title}}": safe(data.beneficiary1Title),
            "{{beneficiary1LastName}}": safe(data.beneficiary1LastName),
            "{{beneficiary1FirstName}}": safe(data.beneficiary1FirstName),
            "{{beneficiary1BirthDate}}": safe(data.beneficiary1BirthDate),
            "{{beneficiary2Title}}": safe(data.beneficiary2Title),
            "{{beneficiary2LastName}}": safe(data.beneficiary2LastName),
            "{{beneficiary2FirstName}}": safe(data.beneficiary2FirstName),
            "{{beneficiary2BirthDate}}": safe(data.beneficiary2BirthDate),
            "{{socialSecurityNumber}}": safe(data.socialSecurityNumber),
            "{{addressLine1}}": safe(data.addressLine1),
            "{{addressLine2}}": safe(data.addressLine2),
            "{{postalCode}}": safe(data.postalCode),
            "{{floor}}": safe(data.floor),
            "{{mobilePhone}}": safe(data.mobilePhone),
            "{{homePhone}}": safe(data.homePhone),
            "{{email}}": safe(data.email),
            "{{hasLegalRepresentative}}": yes_no(data.hasLegalRepresentative),
            "{{legalRepLastName}}": safe(data.legalRepLastName),
            "{{legalRepFirstName}}": safe(data.legalRepFirstName),
            "{{legalRepAddressLine1}}": safe(data.legalRepAddressLine1),
            "{{legalRepAddressLine2}}": safe(data.legalRepAddressLine2),
            "{{legalRepPhone}}": safe(data.legalRepPhone),
            "{{legalRepEmail}}": safe(data.legalRepEmail),
            "{{emergency1LastName}}": safe(data.emergency1LastName),
            "{{emergency1FirstName}}": safe(data.emergency1FirstName),
            "{{emergency1AddressLine1}}": safe(data.emergency1AddressLine1),
            "{{emergency1AddressLine2}}": safe(data.emergency1AddressLine2),
            "{{emergency1Phone}}": safe(data.emergency1Phone),
            "{{emergency1Email}}": safe(data.emergency1Email),
            "{{emergency1Relation}}": safe(data.emergency1Relation),
            "{{emergency2LastName}}": safe(data.emergency2LastName),
            "{{emergency2FirstName}}": safe(data.emergency2FirstName),
            "{{emergency2AddressLine1}}": safe(data.emergency2AddressLine1),
            "{{emergency2AddressLine2}}": safe(data.emergency2AddressLine2),
            "{{emergency2Phone}}": safe(data.emergency2Phone),
            "{{emergency2Email}}": safe(data.emergency2Email),
            "{{emergency2Relation}}": safe(data.emergency2Relation),
            "{{mobilityType}}": safe(data.mobilityType),
            "{{aidWalker}}": yes_no(data.aidWalker),
            "{{aidTransferChair}}": yes_no(data.aidTransferChair),
            "{{aidTripodCane}}": yes_no(data.aidTripodCane),
            "{{aidQuadripodCane}}": yes_no(data.aidQuadripodCane),
            "{{aidSimpleCane}}": yes_no(data.aidSimpleCane),
            "{{aidCrutch}}": yes_no(data.aidCrutch),
            "{{docResidenceProof}}": yes_no(data.docResidenceProof),
            "{{docIdentityCard}}": yes_no(data.docIdentityCard),
            "{{docVitaleCard}}": yes_no(data.docVitaleCard),
            "{{docRetirementOrASPA}}": yes_no(data.docRetirementOrASPA),
            "{{docAPA}}": yes_no(data.docAPA),
            "{{docPCH}}": yes_no(data.docPCH),
            "{{docCMIPriorityOrInvalidity}}": yes_no(data.docCMIPriorityOrInvalidity),
            "{{docMedicalCertificate}}": yes_no(data.docMedicalCertificate),
            "{{engagementTransportRules}}": yes_no(data.engagementTransportRules),
            "{{attestAccuracy}}": yes_no(data.attestAccuracy),
            "{{signatureDate}}": safe(data.signatureDate),
            "{{signature}}": safe(data.beneficiary1LastName),
            "{{additionalNotes}}": safe(data.additionalNotes),

            # Aliases courts (faciles a taper dans Word)
            "{{nom}}": safe(data.beneficiary1LastName),
            "{{prenom}}": safe(data.beneficiary1FirstName),
            "{{civilite}}": safe(data.beneficiary1Title),
            "{{date_naissance}}": safe(data.beneficiary1BirthDate),
            "{{nom2}}": safe(data.beneficiary2LastName),
            "{{prenom2}}": safe(data.beneficiary2FirstName),
            "{{civilite2}}": safe(data.beneficiary2Title),
            "{{date_naissance2}}": safe(data.beneficiary2BirthDate),
            "{{secu}}": safe(data.socialSecurityNumber),
            "{{adresse}}": safe(data.addressLine1),
            "{{adresse2}}": safe(data.addressLine2),
            "{{code_postal}}": safe(data.postalCode),
            "{{etage}}": safe(data.floor),
            "{{portable}}": safe(data.mobilePhone),
            "{{telephone}}": safe(data.homePhone),
            "{{mail}}": safe(data.email),
            "{{rep_nom}}": safe(data.legalRepLastName),
            "{{rep_prenom}}": safe(data.legalRepFirstName),
            "{{rep_adresse}}": safe(f"{data.legalRepAddressLine1} {data.legalRepAddressLine2}".strip()),
            "{{rep_tel}}": safe(data.legalRepPhone),
            "{{rep_mail}}": safe(data.legalRepEmail),
            "{{urgence1_nom}}": safe(data.emergency1LastName),
            "{{urgence1_prenom}}": safe(data.emergency1FirstName),
            "{{urgence1_tel}}": safe(data.emergency1Phone),
            "{{urgence2_nom}}": safe(data.emergency2LastName),
            "{{urgence2_prenom}}": safe(data.emergency2FirstName),
            "{{urgence2_tel}}": safe(data.emergency2Phone),
            "{{mobilite}}": safe(data.mobilityType),
            "{{engagement}}": yes_no(data.engagementTransportRules),
            "{{exactitude}}": yes_no(data.attestAccuracy),
            "{{fait_le}}": safe(data.signatureDate),
            "{{infos}}": safe(data.additionalNotes),
        }

        if TEMPLATE_PATH.exists():
            doc = Document(str(TEMPLATE_PATH))
            replaced_count = replace_in_doc(doc, replacements)
            if replaced_count == 0:
                # If no placeholders exist in the template, keep the exact form and append filled values.
                doc.add_page_break()
                doc.add_heading("DONNEES REMPLIES AUTOMATIQUEMENT", level=1)
                for key, value in replacements.items():
                    add_line(doc, key.replace("{{", "").replace("}}", ""), value)
        else:
            doc = Document()
            doc.add_heading("FORMULAIRE D'INSCRIPTION - LA REINETTE", level=1)
            doc.add_paragraph("Template introuvable, document genere sans mise en page modele.")
            for key, value in replacements.items():
                add_line(doc, key.replace("{{", "").replace("}}", ""), value)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=Fiche-Inscription-Transport-Seniors-BLR.docx"
            },
        )
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Generation DOCX impossible: {exc}") from exc
